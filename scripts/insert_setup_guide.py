"""
Inserts a beginner-friendly "Getting Started — Installing & Running the Tool
on Your Computer" section into the EXISTING docs/aCRF_Annotation_Engine_User_Guide.docx,
without regenerating or otherwise altering the rest of that document.

The User Guide is currently hand-maintained (edited/uploaded directly, not
solely produced by build_user_guide.py), so this script builds the new
section in a scratch document using matching styles, then splices its
paragraphs/tables into the real document immediately after the existing
"What it does" intro and before "1. How it works" — i.e. at the very
beginning of the guide, right after the title.

Re-running this script is safe only if the anchor heading text below still
matches the document; it will raise a clear error otherwise rather than
inserting in the wrong place or duplicating the section.
"""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

_HERE = Path(__file__).resolve().parent.parent
TARGET = _HERE / "docs" / "aCRF_Annotation_Engine_User_Guide.docx"

AZ_PURPLE = RGBColor(0x6B, 0x2D, 0x88)
AZ_MAGENTA = RGBColor(0x83, 0x00, 0x51)
GREY = RGBColor(0x6B, 0x6B, 0x6B)
DARK = RGBColor(0x1A, 0x1A, 0x1A)

PROJECT_PATH = r"C:\Users\kjvd631\OneDrive - AZCollaboration\Desktop\acrf_tool"

# Anchor: insert the new section immediately before this existing heading.
ANCHOR_HEADING_TEXT = "1. How it works"
# Guard: refuse to run twice against the same file.
ALREADY_DONE_MARKER = "Getting Started — Installing & Running the Tool on Your Computer"


def _build_scratch_doc() -> Document:
    """Build the new section's content in a fresh Document with styles that
    match the target guide (same AZ colours/sizes), so it blends in exactly."""
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = DARK

    for lvl, sz, col in [("Heading 1", 15, AZ_MAGENTA), ("Heading 2", 12, AZ_PURPLE)]:
        st = doc.styles[lvl]
        st.font.name = "Calibri"
        st.font.size = Pt(sz)
        st.font.color.rgb = col
        st.font.bold = True

    def bullet(text, bold_prefix=None):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        if bold_prefix:
            r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text)
        return p

    def numbered(text):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(text)
        return p

    def note(text, label="Tip: "):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.15)
        r = p.add_run(label); r.bold = True; r.italic = True
        r.font.size = Pt(9.5); r.font.color.rgb = AZ_PURPLE
        r2 = p.add_run(text); r2.italic = True
        r2.font.size = Pt(9.5); r2.font.color.rgb = GREY
        return p

    def cmd(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.2)
        r = p.add_run(text)
        r.font.name = "Consolas"; r.font.size = Pt(10); r.font.color.rgb = AZ_PURPLE; r.bold = True
        return p

    def body(text, space=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space)
        p.add_run(text)
        return p

    def table(headers, rows, widths=None):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for c, txt in zip(t.rows[0].cells, headers):
            c.paragraphs[0].add_run(txt).bold = True
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                run = cells[i].paragraphs[0].add_run(val)
                if i == 0:
                    run.bold = True
        if widths:
            for r in t.rows:
                for i, w in enumerate(widths):
                    r.cells[i].width = Inches(w)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)
        return t

    # ==================== CONTENT ====================
    doc.add_heading(ALREADY_DONE_MARKER, level=1)
    body("This section is written for someone who has never used these programs "
         "before — every step is spelled out, including things that may feel "
         "obvious if you've done this before. You only need to do this once per "
         "computer. After that, starting the tool again takes four short steps "
         "(see “Starting it again next time” near the end of this section). "
         "Follow the steps in order, from top to bottom.", 8)

    doc.add_heading("Words you'll see in this guide", level=2)
    table(
        ["Word", "What it means"],
        [
            ("Folder", "A place on your computer that holds files — like a "
             "physical folder holding papers. The whole tool lives in one folder."),
            ("VS Code", "A free Microsoft program for opening a project folder "
             "and running it. Think of it as Notepad, a file browser, and a "
             "type-a-command box, all in one window."),
            ("Terminal", "A black or dark-blue box inside VS Code where you type "
             "a short instruction (a “command”) and press Enter, instead of "
             "clicking buttons. It's a direct, typed conversation with your "
             "computer."),
            ("Command", "One line of text you type into the terminal and run by "
             "pressing Enter. Every grey/purple box in this section is a command."),
            ("Python", "A programming language. The part of this tool that reads "
             "the CRF PDF and works out the SDTM mapping is written in Python."),
            ("Node.js", "A second programming language / engine. The part of "
             "this tool you see and click in your web browser is built with it."),
            ("Virtual environment (“venv”)", "A private, separate toolbox of "
             "Python add-ons used only by this project, so it can't clash with "
             "anything else on your computer. You create it once."),
            ("Package", "A ready-made piece of code someone else already wrote "
             "that this tool needs (e.g. the piece that reads PDF files). "
             "“Installing packages” just downloads these automatically."),
            ("node_modules", "A folder that gets created automatically the first "
             "time you set up the website part. It holds all its downloaded "
             "packages. You never need to open it."),
            ("localhost", "A web address that only works on your own computer — "
             "used to view the tool in your browser once it's running, like a "
             "private website just for you."),
        ],
        widths=[1.7, 4.5],
    )

    doc.add_heading("Before you start", level=2)
    bullet("someone needs to share the project folder with you. It will appear "
           "automatically in your own OneDrive once shared — you don't need to "
           "download or copy anything by hand. The folder path is:",
           "Access to the project folder — ")
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.2)
    r = p.add_run(PROJECT_PATH)
    r.font.name = "Consolas"; r.font.size = Pt(9.5); r.font.color.rgb = AZ_PURPLE; r.bold = True
    p.paragraph_format.space_after = Pt(8)
    bullet("VS Code, Python, and Node.js are installed via your AZ Software "
           "Store (also called Company Portal / Software Center — the app you "
           "already use to install approved company software). If you don't "
           "have permission to install software, raise an IT ticket asking for "
           "these three by name.", "A Windows laptop with the AZ Software Store — ")
    bullet("about 15–20 minutes, done once. Every step after this is a "
           "copy-paste of a command shown below.", "Time — ")

    doc.add_heading("Step 1 — Get access to the project folder", level=2)
    numbered("Ask the person handing this tool over to you to share the folder "
             "with your AZ account. Once shared, it appears by itself in your "
             "OneDrive — no action needed from you beyond waiting for it to sync.")
    numbered("Open File Explorer (the folder icon on your taskbar) and go to "
             "OneDrive - AZCollaboration → Desktop → acrf_tool. You should see "
             "folders inside such as app, frontend, src, and cache.")
    note("Before continuing, check the little icon on the acrf_tool folder in "
         "File Explorer. A green tick means OneDrive has finished downloading "
         "it; a blue spinning-arrows icon means it's still syncing — wait for "
         "the green tick.")

    doc.add_heading("Step 2 — Install Visual Studio Code (VS Code)", level=2)
    numbered("Open the AZ Software Store from your Start menu.")
    numbered("Search for “Visual Studio Code” and click Install. Wait for it "
             "to finish.")
    numbered("Open VS Code from your Start menu — its icon is a dark blue angle "
             "bracket (< >).")
    note("If VS Code is already on your laptop, skip straight to Step 5.")

    doc.add_heading("Step 3 — Install Python", level=2)
    numbered("In the AZ Software Store, search for “Python”.")
    numbered("Install Python 3.11 (or the newest 3.x version offered — avoid "
             "anything older than 3.10).")
    note("Not available in the store? Raise an IT ticket asking for "
         "“Python 3.11, added to PATH” to be installed on your laptop.")

    doc.add_heading("Step 4 — Install Node.js", level=2)
    numbered("In the AZ Software Store, search for “Node.js”.")
    numbered("Install the version labelled LTS (Long Term Support) — this is "
             "the stable, recommended one.")
    note("Not available in the store? Raise the same kind of IT ticket asking "
         "for “Node.js LTS” to be installed on your laptop.")

    doc.add_heading("Step 5 — Open the project folder in VS Code", level=2)
    numbered("Open VS Code.")
    numbered("Click File (top-left corner) → Open Folder…")
    numbered("Browse to and select the acrf_tool folder from Step 1, then click "
             "Select Folder.")
    body("VS Code will reload and show the project's files and folders down "
         "the left-hand side — this list is called the Explorer panel.", 6)

    doc.add_heading("Step 6 — Open a terminal inside VS Code", level=2)
    body("From here on, every step is done by typing a command shown below "
         "into a terminal and pressing Enter — nothing needs to be clicked or "
         "configured.", 4)
    numbered("In VS Code's top menu, click Terminal → New Terminal.")
    numbered("A dark panel opens at the bottom of the window with a blinking "
             "cursor — this is your terminal. It already points at your project "
             "folder, so there's nothing to navigate to.")
    note("Type each command below exactly as written into this panel, press "
         "Enter, and wait for it to finish (the cursor becomes free again) "
         "before typing the next one.")

    doc.add_heading("Step 7 — Create the Python toolbox (virtual environment)", level=2)
    body("This creates the private, isolated set of Python add-ons for this "
         "project only (see “Virtual environment” above). Run:", 4)
    cmd("python -m venv venv")
    body("Then switch into it. VS Code's terminal is usually PowerShell — if "
         "so, use this line:", 4)
    cmd(".\\venv\\Scripts\\Activate.ps1")
    body("If your terminal instead says Command Prompt at the top, use this "
         "line instead:", 4)
    cmd("venv\\Scripts\\activate.bat")
    note("You'll know it worked when you see (venv) appear at the start of the "
         "line in your terminal. You only create the venv once — every time you "
         "come back later you just re-run the activate line, not the create line.")

    doc.add_heading("Step 8 — Install the Python packages", level=2)
    cmd("pip install -r requirements.txt")
    body("This reads a shopping list of everything the engine needs and "
         "downloads it automatically. It takes a few minutes and shows a lot "
         "of scrolling text — that's normal, just wait for it to finish.", 6)

    doc.add_heading("Step 9 — Start the engine (backend)", level=2)
    cmd("python -m uvicorn app.main:app --reload --port 8000")
    body("Leave this terminal open and running — closing it stops the tool. "
         "You'll know it worked when the last line says something like "
         "“Application startup complete”.", 6)

    doc.add_heading("Step 10 — Open a second terminal for the website (frontend)", level=2)
    body("The engine (Step 9) and the website need to run at the same time, "
         "in two separate terminals side by side.", 4)
    numbered("Click the “+” icon in the terminal panel to open a second, new "
             "terminal (don't close the first one).")
    numbered("In the new terminal, move into the website's folder:")
    cmd("cd frontend")

    doc.add_heading("Step 11 — Install the website's packages (one-time)", level=2)
    cmd("npm install")
    body("This downloads everything the website screen needs into a folder "
         "called node_modules. The first time, it can take a few minutes and "
         "prints a lot of progress text — that's normal. You won't need to "
         "repeat this step unless you're told the project's packages have "
         "changed.", 6)

    doc.add_heading("Step 12 — Start the website", level=2)
    cmd("npm run dev")
    body("After a few seconds you'll see a line similar to: "
         "Local:  http://localhost:5173/", 6)

    doc.add_heading("Step 13 — Open the tool in your browser", level=2)
    numbered("Hold Ctrl and click the http://localhost:5173/ link shown in the "
             "terminal (or copy it into Chrome / Edge).")
    numbered("The aCRF Annotation Engine screen loads. You're ready to upload a "
             "CRF PDF — continue to the next section below.")

    doc.add_heading("Stopping the tool", level=2)
    body("Click into each terminal and press Ctrl + C, then close VS Code. "
         "Nothing is lost — annotated PDFs and job history stay saved on the "
         "computer.", 6)

    doc.add_heading("Starting it again next time (quick version)", level=2)
    body("Steps 1–4, 7 (create), 8, and 11 are one-time setup. Every time "
         "after that, you only need:", 4)
    numbered("Open VS Code → File → Open Folder → the same acrf_tool folder.")
    numbered("Terminal 1 — activate and start the engine:")
    cmd(".\\venv\\Scripts\\Activate.ps1")
    cmd("python -m uvicorn app.main:app --reload --port 8000")
    numbered("Terminal 2 (click “+” for a new one) — start the website:")
    cmd("cd frontend")
    cmd("npm run dev")
    numbered("Open http://localhost:5173 in your browser.")

    doc.add_heading("If something goes wrong", level=2)
    table(
        ["What you see", "What to do"],
        [
            ("“python is not recognized”",
             "Python wasn't added to PATH during install. Reinstall Python and "
             "tick “Add python.exe to PATH”, or ask IT to do this for you."),
            ("“npm is not recognized”",
             "Node.js isn't installed yet, or VS Code needs restarting after "
             "installing it — close VS Code fully and reopen it."),
            ("“running scripts is disabled on this system” "
             "(when activating the venv)",
             "Type Set-ExecutionPolicy -Scope Process -ExecutionPolicy Bypass, "
             "press Enter, then try the activate line again."),
            ("Port 8000 or 5173 already in use",
             "The tool is probably already running in another terminal or "
             "window. Close old terminals, or restart the computer, then try "
             "again."),
            ("Browser shows nothing at localhost:5173",
             "Check both terminals are still open with no red error text — "
             "one must be running the engine (Step 9), the other the website "
             "(Step 12)."),
            ("The acrf_tool folder looks empty or incomplete",
             "OneDrive hasn't finished syncing. Right-click the folder in "
             "File Explorer and choose “Always keep on this device”, then wait."),
        ],
        widths=[2.3, 3.9],
    )

    doc.add_page_break()
    return doc


def main():
    target = Document(TARGET)

    # Guard against double-insertion.
    existing_headings = {p.text.strip() for p in target.paragraphs}
    if ALREADY_DONE_MARKER in existing_headings:
        print("Setup guide already present — no changes made.")
        return

    anchor = None
    for p in target.paragraphs:
        if p.text.strip() == ANCHOR_HEADING_TEXT:
            anchor = p
            break
    if anchor is None:
        raise SystemExit(
            f"Could not find anchor heading {ANCHOR_HEADING_TEXT!r} in "
            f"{TARGET} — refusing to guess an insertion point."
        )

    scratch = _build_scratch_doc()
    anchor_el = anchor._p

    # scratch body's children, excluding the trailing sectPr element.
    scratch_body = scratch.element.body
    scratch_children = [el for el in scratch_body if el.tag != qn("w:sectPr")]

    for el in scratch_children:
        anchor_el.addprevious(deepcopy(el))

    target.save(TARGET)
    print("Inserted Getting Started section into:", TARGET)


if __name__ == "__main__":
    main()
