"""
Generates the aCRF Annotation Engine — Knowledge Transfer (KT) document (.docx).

Audience: senior SDTM / SAS statistical programmers taking ownership of the
tool. Detailed, technical, and framed against the manual aCRF annotation
workflow a stat programmer already knows. Embeds live application screenshots.
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

_HERE = Path(__file__).resolve().parent.parent
SHOTS = _HERE / "docs" / "images"
OUT = _HERE / "docs" / "aCRF_Annotation_Engine_KT.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

AZ_PURPLE = RGBColor(0x6B, 0x2D, 0x88)
AZ_MAGENTA = RGBColor(0x83, 0x00, 0x51)
GREY = RGBColor(0x6B, 0x6B, 0x6B)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
CODEBG = RGBColor(0x2B, 0x2B, 0x2B)
CODEFG = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK

for lvl, sz, col in [("Heading 1", 15, AZ_MAGENTA), ("Heading 2", 12, AZ_PURPLE),
                     ("Heading 3", 11, AZ_PURPLE)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True

for s in doc.sections:
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)


def caption(text):
    p = doc.add_paragraph()
    r = p.add_run("Figure: " + text)
    r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)


def image(name, width=6.1):
    f = SHOTS / name
    if f.exists():
        doc.add_picture(str(f), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
    p.add_run(text)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(2)
    p.add_run(text)
    return p


def body(text, space=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(text)
    r.italic = italic
    return p


def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    r = p.add_run(text)
    r.font.name = "Consolas"; r.font.size = Pt(9); r.font.color.rgb = AZ_PURPLE
    return p


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, txt in zip(t.rows[0].cells, headers):
        c.paragraphs[0].add_run(txt).bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(val)
            if i == 0:
                run.bold = True
    if widths:
        for r in t.rows:
            for i, w in enumerate(widths):
                r.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ============================ TITLE ============================
t = doc.add_paragraph()
r = t.add_run("aCRF Annotation Engine")
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = AZ_MAGENTA
sub = doc.add_paragraph()
r = sub.add_run("Knowledge Transfer  ·  Technical Handover for SDTM Programmers")
r.font.size = Pt(11.5); r.font.color.rgb = GREY
sub.paragraph_format.space_after = Pt(2)
meta = doc.add_paragraph()
r = meta.add_run("Respiratory & Immunology  ·  SDTM-MSG v2.0  ·  Internal R&I Tool")
r.font.size = Pt(9.5); r.font.color.rgb = GREY
meta.paragraph_format.space_after = Pt(12)

body("Purpose of this document. This is a complete technical knowledge transfer "
     "for the aCRF Annotation Engine, written for statistical programmers who "
     "annotate CRFs by hand today. It explains what the tool does, the SDTM "
     "mapping logic behind it (in the detail you would scrutinise a spec with), "
     "the architecture, the knowledge base it draws on, how it is validated, how "
     "to run and maintain it, and where the boundaries and next steps are. No web "
     "or Python background is assumed — where a software concept appears, it is "
     "tied back to the equivalent in a SAS / SDTM workflow.", 10)

# ============================ 0. ONE-PARAGRAPH SUMMARY ============================
doc.add_heading("1.  What the tool is, in one paragraph", level=1)
body("Annotating a blank CRF is normally a manual job: a programmer opens the "
     "blank CRF PDF, and for every collected field decides the SDTM domain, "
     "variable, codelist and any qualifier, then draws a text box on the page in "
     "Adobe Acrobat. For a full study that is several days of careful, repetitive "
     "work. The aCRF Annotation Engine does that first pass automatically. You "
     "give it the blank CRF PDF; it reads every page, identifies the real "
     "collected fields, maps each one to its SDTM domain and variable following "
     "CDISC SDTM-MSG v2.0, and writes a colour-coded annotated PDF with real, "
     "editable annotation boxes plus a CSV mapping spec. A multi-hundred-page CRF "
     "is annotated in under two minutes. You then review, correct anything you "
     "disagree with, and regenerate — so the programmer stays in control of the "
     "final, submission-quality output.", 6)

table(
    ["Manual workflow (today)", "With the engine"],
    [
        ("Programmer reads each field, recalls the mapping", "Engine proposes the mapping with a confidence score"),
        ("Draws each box by hand in Acrobat", "Boxes placed automatically, aligned to each field row"),
        ("Days per study; error-prone, hard to QC", "Minutes per study; every decision logged to a CSV"),
        ("Consistency depends on the individual", "Same rules + standards applied identically every time"),
        ("No reuse between studies", "A learned-mappings table reuses prior annotated CRFs"),
    ],
    widths=[3.0, 3.4],
)

# ============================ 2. INPUTS / OUTPUTS ============================
doc.add_heading("2.  Inputs and outputs", level=1)
doc.add_heading("2.1  What you feed in", level=2)
bullet("the standard numbered AZ-EDC export. Every field carries a 1–3 digit "
       "number used as a parsing anchor.", "Blank / production CRF ")
bullet("position-based layouts where fields are not numbered; positions are "
       "recovered from page geometry.", "Annotated-style / RSG CRF ")
bullet("spec-table style exports (Field Name / Data Type / SAS Label). EDC "
       "scaffolding pages are detected and skipped.", "DB / Raw dataset CRF ")
body("PDF only, up to 150 MB. No pre-cleaning needed — headers, visit folders, "
     "EDC variable definitions and instructions are filtered out automatically.", 6)

doc.add_heading("2.2  What you get back", level=2)
bullet("colour-coded by domain (SDTM-MSG v2.0 colour sequence), with a coloured "
       "dataset header box per form. Boxes are real PDF FreeText annotations — in "
       "Acrobat you can click, drag, resize and edit them exactly like a "
       "hand-annotated CRF. Derived / dictionary-coded variables carry a dashed "
       "border (Origin = Assigned); genuinely non-submitted fields are marked "
       "[NOT SUBMITTED].", "Annotated PDF ")
bullet("one row per field: form, label, annotation, domain, variable, codelist, "
       "confidence, and which resolution tier decided it. This is your "
       "traceability / review artefact and feeds downstream specs.",
       "CSV mapping spec ")
bullet("resolution rate, annotation count, fields extracted vs dropped as noise, "
       "and a tier breakdown so you know how much of the output to trust.",
       "Run statistics ")

# ============================ 3. THE PIPELINE ============================
doc.add_heading("3.  How a CRF flows through the engine", level=1)
body("Internally the run is a fixed pipeline (see app/services.py → run_pipeline). "
     "Each stage is deterministic and inspectable — there is no black box in the "
     "core path. The optional LLM step is the only non-deterministic component and "
     "is off unless an API key is configured.", 6)

table(
    ["Stage", "What happens"],
    [
        ("1. Parse the PDF",
         "Read every page; extract text with x/y positions; parse headers for form "
         "code + visit; identify fields by the number-anchor pattern "
         "(label → format hint → field number → response options)."),
        ("1.5 Study context",
         "From the set of form codes, detect therapeutic-area context (e.g. "
         "oncology) so domain inference is study-aware."),
        ("2. Noise filter",
         "Drop non-data text that leaked through parsing — headers, instructions, "
         "EDC plumbing. Fields that pass WILL be annotated or left explicitly "
         "unresolved; dictionary fields (MedDRA / WHODrug) are KEPT."),
        ("2.5 De-duplicate",
         "Resolve each unique (form, label) only once — a form repeated across "
         "many visit folders is mapped one time, then the result is reused on "
         "every page, so counts are not inflated."),
        ("3. Resolve to SDTM",
         "The heart of the tool. Each unique field is run through the resolution "
         "tiers (Section 4) until one returns a mapping or it is left unresolved."),
        ("4. Write the PDF",
         "Place a colour-coded FreeText box on each field's row, with overlap "
         "avoidance and dashed borders for derived variables; add domain header "
         "boxes and cross-page 'continued' / 'see page' references."),
        ("4.5 Write the CSV",
         "Emit the full field-level mapping spec for traceability and human review."),
    ],
    widths=[1.7, 4.7],
)

# ============================ 4. RESOLUTION LOGIC ============================
doc.add_heading("4.  The SDTM mapping logic (the part that matters most)", level=1)
body("This is where the engine reproduces the judgement a programmer applies by "
     "hand. It is a tiered cascade: cheap, exact, high-confidence rules are tried "
     "first; only what they cannot answer falls through to broader, lower-"
     "confidence strategies. Every result carries the tier that produced it and a "
     "confidence score, both of which appear in the UI and the CSV so you can "
     "target review at the weakest mappings.", 6)

doc.add_heading("4.1  Tier 1 — NOT SUBMITTED (runs first)", level=2)
body("A deterministic classifier for fields that genuinely do not go into any "
     "submission dataset — RSG/internal fields, calculation-only fields, explicit "
     "'do not submit' markers, and CRF workflow instructions. These are matched "
     "by pattern and regex, never guessed. Important nuance fixed recently: "
     "dictionary-coded fields (MedDRA PT/LLT/SOC, WHODrug preferred name, ATC "
     "class, dictionary versions) are NOT treated as NOT SUBMITTED — they have "
     "real SDTM destinations and are handled by Tier 0 (see Section 4.5).", 6)

doc.add_heading("4.2  Tier 0 — deterministic rules, standards & spec (the workhorse)", level=2)
body("Tier 0 carries the great majority of mappings. It is itself layered, and "
     "the layer that matched sets the confidence band:", 6)
table(
    ["Layer", "Source", "Confidence"],
    [
        ("Exact rules", "Form-aware deterministic rules + the learned-mappings "
         "table built from prior annotated CRFs (Section 5).", "≥ 0.98 (High)"),
        ("SDTM Standards", "CDISC SDTM IG standard variable labels, matched by "
         "normalised label.", "≥ 0.92 (Standards)"),
        ("AZ Spec lookup", "AZ RAW→SDTM specification template + corporate "
         "standards, including fuzzy / semantic matching.", "< 0.92 (Lookup)"),
        ("Domain inference", "Study-agnostic fallback: scores the most likely "
         "domain(s) from form code + form name when no direct label match exists.",
         "varies"),
    ],
    widths=[1.5, 4.0, 1.0],
)
body("Tier 0 also enriches a result: it adds supplemental-qualifier mappings "
     "(SUPP--), attaches controlled-terminology codelists, flags derived / "
     "assigned variables (AGE, --DECOD, --DY, …) so the PDF draws them with a "
     "dashed border, and resolves multi-domain fields that map to more than one "
     "variable.", 6)

doc.add_heading("4.3  Tier 3 — optional LLM fallback", level=2)
body("Only the fields that survive every deterministic tier reach Tier 3. If an "
     "Anthropic API key is configured, the field label + form context is sent to "
     "Claude, which proposes an SDTM variable with a rationale; the result is "
     "flagged for human review and never allowed to invent a variable outside the "
     "standard. With no key set, Tier 3 is a no-op and the field is simply left "
     "Unresolved for a programmer to handle. The core tool ships fully functional "
     "without any LLM.", 6)

doc.add_heading("4.4  Findings where-clauses & value-level decode", level=2)
bullet("for Findings domains (VS, LB, EG, …) the engine carries the test code as "
       "a where-clause qualifier — e.g. VSORRES when VSTESTCD = \"WEIGHT\" — exactly "
       "as a hand-annotated CRF does for repeating test panels.",
       "TESTCD where-clauses ")
bullet("controlled-response questions get a value-level decode showing how each "
       "response option maps to its submission value — e.g. for an "
       "'Condition ongoing? [Yes/No]' field mapping to MHENRTPT, the decode is "
       "'ONGOING = Yes'.", "Value-level decode ")
bullet("the usage guard understands that in repeating Findings structures the "
       "same variable (LBORRES) legitimately recurs for every test in a panel — "
       "so it is not falsely de-duplicated.", "Repeating-structure awareness ")

doc.add_heading("4.5  Dictionary-coded fields (MedDRA / WHODrug / ATC)", level=2)
body("Older CRFs display the dictionary-coding columns directly on the form "
     "(Dictionary-Derived Term, Lowest Level Term, Body System or Organ Class, "
     "Medication dictionary text, Preferred Name, ATC code, Active Ingredient, "
     "dictionary versions). Per SDTM-MSG these are not 'unsubmitted' — the coded "
     "variables (--DECOD, --BODSYS, --LLT, …) are submitted with Origin = "
     "Assigned, and dictionary versions go to the SUPP-- qualifiers (MEDDRAV / "
     "WHODRGV). The engine maps each of these to its real variable, derives the "
     "domain from the form (so the same label resolves to AEDECOD on an AE form "
     "and CMDECOD on a CM form), and draws them with the dashed 'derived' border.", 6)

# ============================ 5. KNOWLEDGE BASE ============================
doc.add_heading("5.  The knowledge base the engine reasons over", level=1)
body("None of the mapping logic is hard-coded study knowledge. It is driven by a "
     "set of pre-built lookup files (the cache/ directory), compiled once from the "
     "authoritative source spreadsheets and reference aCRFs. Think of these as the "
     "engine's standards library and its memory of past work.", 6)
table(
    ["Knowledge file", "What it holds (and the source)"],
    [
        ("sdtm_spec_by_dataset.json", "CDISC SDTM IG standard variables and labels "
         "per dataset — the baseline standard."),
        ("az_spec_lookup.json", "AZ RAW→SDTM Specification Template — the "
         "company's study-build mappings."),
        ("ct_lookup.json", "Controlled Terminology (codelists) used to attach the "
         "right codelist to a mapped variable."),
        ("corporate_variables.json", "AZ Corporate SDTM Standards variables / "
         "datasets."),
        ("map_rules.json", "AZ Map Rule spec — deterministic mapping rules."),
        ("learned_mappings.json", "~7,600 label→variable mappings extracted from "
         "~20 already-annotated reference aCRFs. This is the engine's reuse "
         "memory and the reason reproduction accuracy is high."),
        ("sdtm_not_submitted_labels.json", "Dictionary-derived label reference "
         "(now mapped to real variables, not NOT SUBMITTED)."),
        ("cdisc_index.faiss / az_spec_index.faiss", "Vector indexes for semantic "
         "(meaning-based) lookup when an exact label match is not found."),
    ],
    widths=[2.4, 4.0],
)
body("These are rebuilt with scripts/build_cache.py whenever the source "
     "spreadsheets or reference aCRFs change. The learned-mappings table is the "
     "one most worth growing over time (Section 9).", 6)

# ============================ 6. ACCURACY / VALIDATION ============================
doc.add_heading("6.  How well it works, and how that is measured", level=1)
body("Accuracy is measured against a ground-truth set of 5,073 fields taken from "
     "reference annotated CRFs, using scripts/accuracy_report.py. Two numbers are "
     "reported because they answer two different questions:", 6)
table(
    ["Metric", "Meaning"],
    [
        ("≈ 96% reproduction (learned table ON)", "On CRFs similar to ones it has "
         "seen, with the learned-mappings memory enabled, the engine reproduces "
         "the human annotation for ~96% of fields. This is the day-to-day mode."),
        ("≈ 8% generalisation (learned table OFF)", "With the reuse memory "
         "disabled — pure rules + standards on unseen fields — exact "
         "domain+variable match. This is the hard floor on a brand-new CRF and is "
         "the honest measure of out-of-the-box generalisation."),
    ],
    widths=[2.6, 3.8],
)
body("The gap between the two is exactly why the learned table and the "
     "human-in-the-loop review matter: every correction a programmer makes can be "
     "fed back to raise the floor over time. Confidence scores should be "
     "interpreted as review priority, not as a guarantee — the 'All Confidence' "
     "filter in the UI exists so you can focus review on the weakest mappings "
     "first.", 6)

# ============================ 7. THE APPLICATION (UI) ============================
doc.add_heading("7.  Using the application", level=1)
body("The tool ships as a small web application: a browser front-end for upload, "
     "review and correction, talking to a Python service that runs the pipeline. "
     "A programmer never touches code to use it.", 6)
image("01_upload.png", 5.2)
caption("Upload screen — drag in a blank CRF PDF; the engine processes it and "
        "opens the results automatically.")
body("On the job page, four cards summarise quality (resolution rate, annotations "
     "written, fields extracted vs noise, unique forms), a donut shows resolved vs "
     "unresolved, and a tier breakdown shows HOW each mapping was decided.", 4)
image("04_stats_charts.png", 6.0)
caption("Job dashboard — quality metrics, resolution donut, tier breakdown, and "
        "the filterable mappings table.")
body("Every field is listed with its annotation, domain, variable and a confidence "
     "bar, split into Resolved / Unresolved tabs. Filter by form, SDTM class, "
     "confidence threshold, or free text; sort on any column.", 4)
image("05_edit_inline.png", 6.0)
caption("Field-mappings table with one row open for inline editing.")

doc.add_heading("7.1  Human-in-the-loop correction", level=2)
body("This is the key point for a programmer audience: the engine's output is a "
     "starting point you control, not a final answer it imposes.", 4)
numbered("Click the pencil on any row to edit the annotation.")
numbered("Type the corrected mapping (comma-separate multiple annotations on one "
         "field, e.g. VS.VSORRES, VS.VSTESTCD).")
numbered("Press the green check to stage it — staged, not yet committed.")
numbered("Click 'Save & Regenerate PDF' in the sticky footer; the engine rewrites "
         "the annotated PDF with your corrections. 'Discard All' reverts.")
image("06_pending_footer.png", 6.0)
caption("A staged edit highlights the row and raises the Save & Regenerate footer.")
body("Every run is kept in Job History — re-open, re-download the PDF, export the "
     "CSV, or keep editing any previous study.", 4)
image("07_jobs.png", 5.4)
caption("Job History — every annotation run, re-openable and re-downloadable.")
image("08_pdf_crop.png", 5.0)
caption("Output detail — editable, colour-coded FreeText boxes and dashed "
        "derived / NOT SUBMITTED markers, aligned to each CRF field row.")

# ============================ 8. ARCHITECTURE ============================
doc.add_heading("8.  Architecture & technology (for whoever maintains it)", level=1)
body("The stack, in plain terms with the SDTM-programmer equivalent where useful:", 6)
table(
    ["Component", "Technology", "Role"],
    [
        ("PDF engine", "Python + PyMuPDF (fitz)", "Reads the blank CRF and writes "
         "the annotated PDF with real FreeText annotation objects."),
        ("Resolution engine", "Python (src/resolution/*)", "The tiered SDTM "
         "mapping logic — the equivalent of your annotation judgement, codified."),
        ("API service", "FastAPI + uvicorn", "Receives the upload, runs the "
         "pipeline, serves stats / downloads. Like a stored procedure exposed "
         "over the web."),
        ("Front-end", "React + TypeScript (Vite)", "The browser UI — upload, "
         "dashboard, mappings table, inline edit."),
        ("Test suite", "pytest (78 tests)", "Locks in extractor, mapping, "
         "annotation-format and regression behaviour."),
    ],
    widths=[1.5, 1.9, 3.0],
)
doc.add_heading("8.1  Where things live in the repository", level=2)
table(
    ["Path", "Contents"],
    [
        ("src/pdf_parser/", "PDF reading: extractor, field identification "
         "(number-anchor), header parsing, context windows."),
        ("src/resolution/", "The mapping tiers: tier0_rules, tier1_not_submitted, "
         "tier3_llm, domain_inferencer, findings_qualifier, value_decoder, "
         "noise_filter, usage_guard, models."),
        ("src/annotator/", "pdf_writer (box placement, colours, dashed borders, "
         "overlap avoidance) and mapping_export (CSV)."),
        ("src/cache_builder/", "Parsers that compile the source spreadsheets into "
         "the cache/ lookup files."),
        ("app/", "FastAPI service: routers/annotate.py (endpoints), services.py "
         "(run_pipeline), schemas.py."),
        ("frontend/", "React app (pages: Upload, JobDetail, Jobs)."),
        ("cache/", "The pre-built knowledge base (Section 5)."),
        ("scripts/", "build_cache, accuracy_report/metrics, doc builders, "
         "extract_from_acrf (grows the learned table)."),
        ("tests/", "The pytest suite."),
        ("config/settings.py", "All paths, thresholds and tuneables in one place."),
    ],
    widths=[1.8, 4.6],
)

# ============================ 9. OPERATE / MAINTAIN ============================
doc.add_heading("9.  Running, maintaining and extending it", level=1)
doc.add_heading("9.1  Run it", level=2)
body("Command-line, one CRF (no web UI needed):", 2)
code("python run.py input/my_crf.pdf --output output/annotated.pdf")
body("Start the web application (API + served front-end):", 2)
code("uvicorn app.main:app --host 0.0.0.0 --port 8000")
body("Run the test suite / the accuracy harness:", 2)
code("python -m pytest -q\npython scripts/accuracy_report.py")

doc.add_heading("9.2  Maintain the knowledge base", level=2)
bullet("when the AZ spec spreadsheets, controlled terminology, or corporate "
       "standards change, rebuild the lookups.",
       "Rebuild caches — scripts/build_cache.py ")
bullet("point scripts/extract_from_acrf.py at newly annotated reference aCRFs to "
       "add their label→variable mappings to learned_mappings.json. This is the "
       "single highest-leverage maintenance task — it directly raises reproduction "
       "accuracy on future studies.", "Grow the learned table ")
bullet("deterministic rules live in config/rules and src/resolution/tier0_rules. "
       "Add or refine a rule when a class of fields is consistently mis-mapped.",
       "Tune rules ")

doc.add_heading("9.3  Recent engineering work (context for the new owner)", level=2)
bullet("dictionary-coded fields (MedDRA / WHODrug / ATC) now map to their real "
       "SDTM variables and SUPP qualifiers instead of being marked NOT SUBMITTED "
       "(Section 4.5).", "MedDRA / WHODrug mapping fix ")
bullet("on tightly-pitched findings forms, annotation boxes used to cascade "
       "downward off their field rows; box density is now detected by actual row "
       "pitch, eliminating the drift (measured: 354 drifted boxes → 0 on a "
       "reference CRF).", "Annotation drift fix ")
bullet("clearer upload progress (live staged checklist) and friendly handling of "
       "password-protected, scanned/image-only, corrupt, and non-CRF PDFs.",
       "UX & edge cases ")

# ============================ 10. LIMITATIONS ============================
doc.add_heading("10.  Known limitations & boundaries", level=1)
bullet("a brand-new CRF with unfamiliar labels relies on rules + standards only "
       "(~8% exact out-of-the-box); the learned table and human review close the "
       "gap. Treat first-pass output on a novel CRF as a draft.",
       "Generalisation floor ")
bullet("scanned / image-only PDFs have no extractable text; the engine now detects "
       "and rejects these with a clear message rather than producing nothing. (OCR "
       "is a future enhancement.)", "Image-only CRFs ")
bullet("the deterministic core never invents variables. The optional LLM tier can "
       "propose unseen mappings but is off by default and always flagged for "
       "review.", "No silent guessing ")
bullet("confidence is a review-priority signal, not a validated probability — it "
       "should not be used as an automated accept/reject gate without human "
       "sign-off.", "Confidence ≠ guarantee ")

doc.add_heading("10.1  Production-readiness roadmap (forward-looking)", level=2)
body("The mapping engine is mature; the surrounding service is currently a "
     "single-instance prototype. To run it as a shared, validated service the main "
     "items are: persistent storage for jobs and artefacts (today they live in "
     "memory and temp folders), asynchronous job processing (today the pipeline "
     "runs inside one HTTP request), authentication / access control, and an audit "
     "trail of who annotated and edited what (relevant for GxP / 21 CFR Part 11). "
     "These are infrastructure concerns and do not affect the correctness of the "
     "annotation logic described above.", 6)

# ============================ 11. GLOSSARY ============================
doc.add_heading("11.  Glossary — software term ↔ familiar concept", level=1)
table(
    ["Term in this doc", "Think of it as"],
    [
        ("Pipeline", "The ordered set of steps a CRF goes through, like a program "
         "flow."),
        ("Tier (0/1/3)", "A mapping strategy, tried in priority order — exact "
         "rules first, broad/fuzzy last."),
        ("Learned-mappings table", "The engine's memory of how fields were "
         "annotated on past studies, reused on new ones."),
        ("Cache files", "The pre-compiled standards library (SDTM IG, AZ spec, CT) "
         "the engine looks up against."),
        ("FreeText annotation", "A real, clickable text box in the PDF — the same "
         "object you create by hand in Acrobat."),
        ("Derived / dashed border", "Origin = Assigned variables (--DECOD, AGE, "
         "study day) — not collected, shown dashed per SDTM-MSG."),
        ("API / FastAPI", "The service layer that runs the engine; like a stored "
         "procedure called over the web."),
        ("Front-end / React", "The browser screens used to upload and review."),
        ("pytest", "The automated regression test suite that guards behaviour."),
    ],
    widths=[2.2, 4.2],
)

body("Questions on any section can be walked through live against the running "
     "tool and the source files referenced above.", 6, italic=True)

doc.save(str(OUT))
print("Saved:", OUT)
