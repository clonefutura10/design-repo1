"""
Generates the aCRF Annotation Engine — Leadership Overview (Word .docx).

Audience: R&I leadership / SDTM governance. Value, speed, compliance, and risk
framing rather than a button-by-button UI walkthrough. Embeds the same live
screenshots captured from the running application.
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

_HERE = Path(__file__).resolve().parent.parent
SHOTS = _HERE / "docs" / "images"
OUT = _HERE / "docs" / "aCRF_Annotation_Engine_Leadership_Overview.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

AZ_PURPLE = RGBColor(0x6B, 0x2D, 0x88)
AZ_MAGENTA = RGBColor(0x83, 0x00, 0x51)
GREY = RGBColor(0x6B, 0x6B, 0x6B)
DARK = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "Calibri"; normal.font.size = Pt(10.5); normal.font.color.rgb = DARK
for lvl, sz, col in [("Heading 1", 14, AZ_MAGENTA), ("Heading 2", 12, AZ_PURPLE)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"; st.font.size = Pt(sz); st.font.color.rgb = col; st.font.bold = True

for s in doc.sections:
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)


def image(name, width=5.6):
    doc.add_picture(str(SHOTS / name), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        p.add_run(bold_prefix).bold = True
    p.add_run(text)
    return p


def body(text, space=8):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(space)
    p.add_run(text); return p


def kv_table(rows, headers=("Metric", "What it means")):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid Accent 1"; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for c, txt in zip(tbl.rows[0].cells, headers):
        c.paragraphs[0].add_run(txt).bold = True
    for a, b in rows:
        cells = tbl.add_row().cells
        cells[0].paragraphs[0].add_run(a).bold = True
        cells[1].paragraphs[0].add_run(b)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


# ============================ TITLE ============================
t = doc.add_paragraph()
r = t.add_run("aCRF Annotation Engine")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = AZ_MAGENTA
sub = doc.add_paragraph()
r = sub.add_run("Leadership Overview  ·  SDTM Automation for Respiratory & Immunology")
r.font.size = Pt(11); r.font.color.rgb = GREY
sub.paragraph_format.space_after = Pt(10)

# ============================ 1. WHAT IT DOES ============================
doc.add_heading("What it does", level=1)
body("The aCRF Annotation Engine turns a blank study CRF into a fully SDTM-annotated "
     "CRF automatically — work that is otherwise done by hand over several days for "
     "every study. It maps each collected field to its SDTM domain and variable "
     "following CDISC SDTM-MSG v2.0, then produces a submission-ready annotated PDF "
     "with a full field-level traceability export. A programmer reviews and signs "
     "off rather than annotating from scratch.")

# ============================ 2. AT A GLANCE ============================
doc.add_heading("At a glance", level=1)
kv_table([
    ("Speed", "Annotates a ~500-page CRF in under two minutes, versus days of manual work per study."),
    ("Coverage", "~93% of fields auto-mapped to the correct SDTM domain and variable on a representative R&I study."),
    ("Compliance", "Output follows CDISC SDTM-MSG v2.0 annotation conventions, supporting FDA Study Data Technical Conformance Guide expectations for the aCRF."),
    ("Risk control", "Human-in-the-loop by design — every mapping is surfaced for review with a confidence score; nothing reaches the deliverable unvalidated."),
    ("Reproducibility", "Each run is versioned and stamped with its tool/rule provenance, and persisted in a job history for audit."),
], headers=("", ""))

# ============================ 3. HOW IT WORKS ============================
doc.add_heading("How it works", level=1)
bullet("any standard AZ-EDC CRF export (Blank, RSG, or DB format) — no pre-cleaning "
       "required; EDC scaffolding is detected and skipped.", "Upload  ")
bullet("reads every page to identify forms, visits and collected fields, filtering "
       "out non-collected noise.", "Extract  ")
bullet("resolves each field to its SDTM domain, variable, codelist and qualifier, "
       "with a confidence score and supplemental-dataset support.", "Map  ")
bullet("a colour-coded, MSG-compliant annotated PDF plus a CSV traceability file.",
       "Output  ")
image("01_upload.png", 5.0)

# ============================ 4. QUALITY ASSURANCE ============================
doc.add_heading("Quality assurance", level=1)
body("Every run opens on a quality dashboard that makes coverage and trustworthiness "
     "visible at a glance:")
kv_table([
    ("Resolution Rate", "Share of extracted fields mapped to an SDTM variable (or marked NOT SUBMITTED) — the headline coverage number."),
    ("Confidence tiering", "Each mapping is graded — High Confidence (exact rules), SDTM Standards, Standards Lookup, Not Submitted, Unresolved — so review effort targets the weakest mappings, not the whole CRF."),
    ("Annotations / Forms", "Volume placed in the output and the breadth of CRF forms covered."),
])
image("04_stats_charts.png", 6.1)

# ============================ 5. HUMAN REVIEW & OVERRIDE ============================
doc.add_heading("Human review & override", level=1)
body("The engine's output is a starting point, not a locked deliverable. Every mapping "
     "is reviewable in a filterable table split by Resolved / Unresolved; programmers "
     "focus on low-confidence annotations and leave exact, rule-based matches with "
     "minimal scrutiny. Any annotation can be corrected inline and the PDF regenerated "
     "in one click — keeping a qualified programmer accountable for the final result.")
image("06_pending_footer.png", 6.1)

# ============================ 6. OUTPUT & COMPLIANCE ============================
doc.add_heading("Output & compliance", level=1)
body("The deliverable is a fully SDTM-MSG v2.0-compliant annotated PDF with colour-coded "
     "domain headers, editable FreeText annotation boxes (adjustable in Adobe Acrobat) and "
     "[NOT SUBMITTED] markers for non-collected fields. A companion CSV provides field-level "
     "traceability for review sign-off and downstream SDTM spec generation. Because the "
     "annotations are real PDF objects rather than flattened images, the output drops "
     "directly into the existing submission workflow and remains fully adjustable.")
image("08_pdf_crop.png", 5.0)

doc.save(str(OUT))
print("Saved:", OUT)
