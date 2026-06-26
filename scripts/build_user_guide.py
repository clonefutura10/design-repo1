"""
Generates the aCRF Annotation Engine — User Guide (Word .docx).

Audience: SDTM statistical programmers and R&I leadership. Embeds live
screenshots captured from the running application.
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_HERE = Path(__file__).resolve().parent.parent
SHOTS = _HERE / "docs" / "images"
OUT = _HERE / "docs" / "aCRF_Annotation_Engine_User_Guide.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

AZ_PURPLE = RGBColor(0x6B, 0x2D, 0x88)
AZ_MAGENTA = RGBColor(0x83, 0x00, 0x51)
GREY = RGBColor(0x6B, 0x6B, 0x6B)
DARK = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()

# ---- base styles ----
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK

for lvl, sz, col in [("Heading 1", 15, AZ_MAGENTA), ("Heading 2", 12, AZ_PURPLE)]:
    st = doc.styles[lvl]
    st.font.name = "Calibri"
    st.font.size = Pt(sz)
    st.font.color.rgb = col
    st.font.bold = True

# tighten margins for a denser, professional layout
for s in doc.sections:
    s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)


def caption(text):
    p = doc.add_paragraph()
    r = p.add_run("Figure: " + text)
    r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)


def image(name, width=6.1):
    doc.add_picture(str(SHOTS / name), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix); r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def body(text, space=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    p.add_run(text)
    return p


# ============================ TITLE ============================
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = t.add_run("aCRF Annotation Engine")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = AZ_MAGENTA
sub = doc.add_paragraph()
r = sub.add_run("User Guide  ·  SDTM Automation for Respiratory & Immunology")
r.font.size = Pt(11); r.font.color.rgb = GREY
sub.paragraph_format.space_after = Pt(10)

body("This tool takes a blank study CRF (PDF) and automatically produces a fully "
     "annotated CRF (aCRF) in which every collected field is mapped to its SDTM "
     "domain and variable, following CDISC SDTM-MSG v2.0 conventions. It is built "
     "for SDTM programmers: the engine does the first pass in seconds, and you "
     "review, correct, and regenerate from a web interface — no manual box-drawing "
     "in Acrobat. This guide walks through every screen and feature.", 8)

# ============================ 1. INPUTS ============================
doc.add_heading("1.  What you can upload", level=1)
body("Open the Upload & Annotate screen and drag in a single CRF PDF (or click to "
     "browse). PDF only, up to 150 MB. The engine supports the CRF formats issued "
     "by AZ-EDC:")
bullet("— a numbered, anchor-based blank CRF (the standard EDC export).", "Blank / production CRF ")
bullet("— position-based layouts where fields are not numbered; the engine "
       "recovers field positions from the page geometry.", "Annotated-style / RSG CRF ")
bullet("— spec-table (“Field Name / Data Type / SAS Label”) exports. EDC "
       "scaffolding pages are detected and skipped so only the real CRF screens are "
       "annotated.", "DB / Raw dataset CRF ")
body("You do not need to pre-clean the file. Form headers, visit folders, EDC "
     "variable definitions and other non-collected text are filtered out "
     "automatically.", 8)
image("01_upload.png", 5.4)
caption("Upload screen — drag-and-drop with a five-step summary of the pipeline.")

# ============================ 2. DASHBOARD METRICS ============================
doc.add_heading("2.  Reading the results dashboard", level=1)
body("When processing finishes (about 30–90 seconds for a large CRF) you land on "
     "the job page. The four cards at the top give you an at-a-glance quality read:")

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for c, txt in zip(hdr, ["Metric", "What it tells you"]):
    c.paragraphs[0].add_run(txt).bold = True
rows = [
    ("Resolution Rate", "Share of extracted fields the engine mapped to an SDTM variable (or marked NOT SUBMITTED). Your headline coverage number."),
    ("Annotations Written", "Count of annotation boxes placed in the output PDF, and how many pages carry them."),
    ("Fields Extracted", "Total candidate fields read from the CRF, and how many were dropped as noise (EDC plumbing, instructions)."),
    ("Unique Forms", "Number of distinct CRF form types detected across the document."),
]
for a, b in rows:
    cells = tbl.add_row().cells
    cells[0].paragraphs[0].add_run(a).bold = True
    cells[1].paragraphs[0].add_run(b)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

body("Below the cards, two panels explain how that coverage was achieved:")
bullet("the resolved-vs-unresolved split at a glance.", "Resolution Overview (donut) ")
bullet("how each mapping was decided, so you know how much to trust it — "
       "High Confidence (exact rules), SDTM Standards, Standards Lookup, "
       "Not Submitted, and Unresolved. Higher tiers need less review.",
       "Resolution Tier Breakdown ")
image("04_stats_charts.png", 6.1)
caption("Job dashboard — quality metrics, resolution donut, tier breakdown, and the "
        "filter toolbar above the mappings table.")

# ============================ 3. MAPPINGS TABLE ============================
doc.add_heading("3.  The field-mappings table — review & filter", level=1)
body("Every field the engine read is listed here, with its CRF label, the SDTM "
     "annotation, the domain, the variable, and a confidence bar. Two tabs split "
     "the work: Resolved and Unresolved (the fields that need your attention). "
     "Use the toolbar to narrow the list:")
bullet("free-text match on field label, form code, or annotation.", "Search ")
bullet("show one CRF form at a time.", "All Forms ")
bullet("filter by SDTM class (Events, Interventions, Findings, Findings About, "
       "Special Purpose, Relationship).", "All Classes ")
bullet("show only mappings at or above ≥90% / ≥95% / ≥98%, to focus review on the "
       "weaker ones.", "All Confidence ")
body("Every column header is click-to-sort. The footer shows how many fields match "
     "your current filters.", 6)
image("05_edit_inline.png", 6.1)
caption("Mappings table with one row open for editing; NOT SUBMITTED and tier badges "
        "sit under each field label.")

# ============================ 4. EDIT + REGENERATE ============================
doc.add_heading("4.  Correcting an annotation and regenerating the PDF", level=1)
body("You are in control of the final output. To change any mapping:")
p = doc.add_paragraph(style="List Number"); p.add_run("Click the pencil icon on the row. The annotation becomes an editable field.")
p = doc.add_paragraph(style="List Number"); p.add_run("Type the corrected annotation. For multiple annotations on one field, separate them with commas (e.g. VS.VSORRES, VS.VSTESTCD).")
p = doc.add_paragraph(style="List Number"); p.add_run("Press the green check (or Enter) to stage the edit. The row is highlighted and counted as a pending edit; nothing is committed yet.")
p = doc.add_paragraph(style="List Number"); p.add_run("When you have staged all your edits, click Save & Regenerate PDF in the sticky footer. The engine rewrites the annotated PDF with your corrections; Discard All reverts pending edits.")
body("You can also copy any annotation to the clipboard with the copy icon — handy "
     "when reusing a mapping across similar fields.", 6)
image("06_pending_footer.png", 6.1)
caption("A staged edit highlights the row and raises the “Save & Regenerate PDF” footer.")

# ============================ 5. OUTPUT PDF ============================
doc.add_heading("5.  The annotated PDF output", level=1)
body("Download the finished aCRF with the Download PDF button (top-right of the job "
     "page). The full field-level mapping can be exported separately as a "
     "traceability spreadsheet with Export CSV — useful for review sign-off and for "
     "feeding downstream specs. The annotated PDF follows SDTM-MSG v2.0:")
bullet("colour-coded by domain, with the dataset name in a coloured header box at "
       "the top-left of each form (e.g. CM, MH).", "Domain headers ")
bullet("each box is filled in the domain's colour with a darker frame; black text "
       "carries the variable name. Derived / dictionary-coded variables and "
       "non-collected fields use a dashed grey [NOT SUBMITTED] box.", "Variable boxes ")
bullet("the boxes are real PDF FreeText annotations, not flattened pixels. In Adobe "
       "Acrobat a programmer can click any box to drag, resize, re-position, or edit "
       "its text, exactly as with a hand-annotated CRF — so the engine's output is a "
       "starting point you stay free to adjust.", "Fully editable in Acrobat ")
image("08_pdf_crop.png", 5.2)
caption("Output detail — editable, colour-coded FreeText annotation boxes and dashed "
        "NOT SUBMITTED markers, aligned to each CRF field row.")

# ============================ 6. JOB HISTORY ============================
doc.add_heading("6.  Job history", level=1)
body("The Job History screen lists every annotation run with its filename, job ID, "
     "annotation count, resolution rate, and status. Click View to reopen any job, "
     "re-download its PDF, or continue editing.", 6)
image("07_jobs.png", 5.6)
caption("Job History — re-open, re-download, or keep editing any previous run.")

doc.save(str(OUT))
print("Saved:", OUT)
